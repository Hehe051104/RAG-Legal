import docx
import re
import json
import os

# 处理法律原文的  优化过
def parse_code_perfect(file_path, output_json):
    doc = docx.Document(file_path)
    legal_records = []

    file_name = os.path.basename(file_path)
    file_name_without_ext = os.path.splitext(file_name)[0]
    # 使用正则表达式处理文件名中的空格、下划线或括号
    split_parts = re.split(r'[_\s\u3000（(]', file_name_without_ext)
    source_name = split_parts[0].strip()

    # 层级
    current_book = ""     # 编
    current_subbook = ""  # 分编
    current_chapter = ""  # 章
    current_section = ""  # 节
    is_inside_article = False # 标记当前是否处于法条内部  用于出来一个法条下的连续段

    # 3. 正则表达式定义
    re_book = re.compile(r'^(第[一二三四五六七八九十百千零]+编)[\s\u3000]+(.+)$')
    re_subbook = re.compile(r'^(第[一二三四五六七八九十百千零]+分编)[\s\u3000]+(.+)$')
    re_chapter = re.compile(r'^(第[一二三四五六七八九十百千零]+章)[\s\u3000]+(.+)$')
    re_section = re.compile(r'^(第[一二三四五六七八九十百千零]+节)[\s\u3000]+(.+)$')
    re_article = re.compile(r'^(第[一二三四五六七八九十百千零]+条(?:之[一二三四五六七八九十]+)?)[\s\u3000]*(.*)$')
    re_appendix = re.compile(r'^附[\s\u3000]*则$')

    # 遍历段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 匹配各级标题 (一旦匹配到标题，必须关闭追加开关)

        # 匹配附则
        if re_appendix.match(text):
            current_book = "附则"
            current_subbook = ""
            current_chapter = ""
            current_section = ""
            is_inside_article = False # 关闭开关
            continue

        # 匹配编
        book_match = re_book.match(text)
        if book_match:
            current_book = text
            current_subbook = ""
            current_chapter = ""
            current_section = ""
            is_inside_article = False # 关闭开关
            continue

        # 匹配分编
        subbook_match = re_subbook.match(text)
        if subbook_match:
            current_subbook = text
            current_chapter = ""
            current_section = ""
            is_inside_article = False # 关闭开关
            continue

        # 匹配章
        chapter_match = re_chapter.match(text)
        if chapter_match:
            current_chapter = text
            current_section = ""
            is_inside_article = False # 关闭开关
            continue

        # 匹配节
        section_match = re_section.match(text)
        if section_match:
            current_section = text
            is_inside_article = False # 关闭开关
            continue

        # 匹配“条” (开启追加开关)
        article_match = re_article.match(text)
        if article_match:
            article_num = article_match.group(1)
            content = article_match.group(2).strip()

            record = {
                "id": f"{source_name}_{article_num}",
                "article_number": article_num,
                "hierarchy": {
                    "book": current_book,
                    "subbook": current_subbook,
                    "chapter": current_chapter,
                    "section": current_section
                },
                "content": content if content else "", # 可能条号后暂时没内容
                "source": source_name,
                "doc_type": "law"
            }
            legal_records.append(record)
            is_inside_article = True # 开启开关：告诉程序接下来的文字都属于这一条
            continue

        # 解决连续的多段
        if is_inside_article and legal_records:
            # 如果原本 content 有内容，加个换行；如果没有（条号占一行），直接赋值
            if legal_records[-1]["content"]:
                legal_records[-1]["content"] += "\n" + text
            else:
                legal_records[-1]["content"] = text

    # 保存结果
    output_dir = os.path.dirname(output_json)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(legal_records, f, ensure_ascii=False, indent=4)

    print(f"解析结束！共提取 {len(legal_records)} 条法律条文。")
    print(f"JSON已保存至：{os.path.abspath(output_json)}")

# parse_code_perfect("法律原文/中华人民共和国刑法_20201226.docx", "code_json/criminal_law_452.json",'Criminal_law')

# 处理司法解释的  相较于法律条文,只需要识别chapter和artical,正则表达式做了改变,其他一样
def parse_interpretation_perfect(file_path, output_json):
    doc = docx.Document(file_path)
    legal_records = []

    # 1. 提取 source_name (按照你之前的 split 逻辑)
    file_name = os.path.basename(file_path)
    file_name_without_ext = os.path.splitext(file_name)[0]
    split_parts = re.split(r'[_\s\u3000（(]', file_name_without_ext)
    source_name = split_parts[0].strip()

    # 2. 层级与状态初始化
    current_chapter = ""
    is_inside_article = False # 追加开关

    # 3. 正则表达式
    # 匹配司法解释的大章节：一、关于一般规定 (注意那个顿号 、)
    re_chapter = re.compile(r'^([一二三四五六七八九十]+、.+)$')
    # 匹配法条：第X条 (后面允许直接跟内容或换行)
    re_article = re.compile(r'^(第[一二三四五六七八九十百千零]+条(?:之[一二三四五六七八九十]+)?)[\s\u3000]*(.*)$')

    # 4. 遍历解析
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 匹配章节标题
        # 匹配到“一、”时，更新章节名，并关闭法条追加开关
        chapter_match = re_chapter.match(text)
        if chapter_match:
            current_chapter = text
            is_inside_article = False
            continue

        # 匹配“条”
        article_match = re_article.match(text)
        if article_match:
            article_num = article_match.group(1)
            content = article_match.group(2).strip()

            # 创建符合你 metadata 要求的 record
            record = {
                "id": f"{source_name}_{article_num}",
                "article_number": article_num,
                "hierarchy": {
                    "book": "",      # 司法解释通常无编
                    "subbook": "",   # 无分编
                    "chapter": current_chapter,
                    "section": ""    # 无节
                },
                "content": content if content else "",
                "source": source_name,
                "doc_type": "interpretation"
            }
            legal_records.append(record)
            is_inside_article = True # 开启追加开关
            continue

        # --- C. 处理子项和多段落补充内容 (如：(一)、(二)) ---
        # 只要开关是开着的，且列表不为空，就把文字粘在上一条的屁股后面
        # 由于 (一) 匹配不上 re_chapter 和 re_article，它会自然落入这里
        if is_inside_article and legal_records:
            if legal_records[-1]["content"]:
                legal_records[-1]["content"] += "\n" + text
            else:
                legal_records[-1]["content"] = text

    # 创建文件夹并保存
    output_dir = os.path.dirname(output_json)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(legal_records, f, ensure_ascii=False, indent=4)

    print(f"司法解释解析结束！提取 {len(legal_records)} 条。")
    print(f"JSON已保存至：{os.path.abspath(output_json)}")


# 处理裁判文书/案例
def parse_case(file_path, output_json):
    doc = docx.Document(file_path)
    legal_records = []

    file_name = os.path.basename(file_path)
    file_name_without_ext = os.path.splitext(file_name)[0]
    # 提取来源名（法院名或案例来源）
    split_parts = re.split(r'[_\s　（(]', file_name_without_ext)
    source_name = split_parts[0].strip()

    # 提取案号（如果文件名中包含案号）
    # 支持格式：指导案例1号、（2023）最高法民申1234号、案例-张某故意杀人案
    case_number_match = re.search(r'指导案例\s*(\d+)\s*号', file_name_without_ext)
    if case_number_match:
        case_number = f"指导案例{case_number_match.group(1)}号"
    else:
        case_number_match = re.search(r'[（(]\d{4}[）)].*?号', file_name_without_ext)
        if case_number_match:
            case_number = case_number_match.group(0)
        else:
            # 从文件名中提取有意义的案例名
            # 去除日期后缀等无关信息
            clean_name = re.sub(r'_\d{8}$', '', file_name_without_ext)
            case_number = clean_name

    # 提取来源名 - 优先使用案例号作为来源
    if case_number.startswith("指导案例"):
        source_name = case_number
    elif "案例-" in source_name or "案例_" in source_name:
        source_name = case_number

    # 裁判文书结构识别正则
    re_court = re.compile(r'.+人民法院')
    re_date = re.compile(r'[一二三四五六七八九十○〇]{4}年[一二三四五六七八九十]{1,2}月[一二三四五六七八九十]{1,2}日')
    re_section_header = re.compile(r'^(本院认为|经审理查明|事实认定|判决如下|裁定如下|本院查明|辩称|诉称|上诉人诉称|被上诉人辩称)')

    # 解析段落
    full_text = []
    court_name = source_name
    judgment_date = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text.append(text)

        # 提取法院名
        court_match = re_court.search(text)
        if court_match and not judgment_date:
            court_name = court_match.group(0)

        # 提取日期
        date_match = re_date.search(text)
        if date_match:
            judgment_date = date_match.group(0)

    # 按段落结构分段
    current_section = "全文"
    sections = {"全文": []}

    for text in full_text:
        section_match = re_section_header.match(text)
        if section_match:
            current_section = section_match.group(1)
            if current_section not in sections:
                sections[current_section] = []
            sections[current_section].append(text)
        else:
            if current_section not in sections:
                sections[current_section] = []
            sections[current_section].append(text)

    total_text = "\n".join(full_text)

    safe_name = re.sub(r'[^\w一-鿿]', '_', file_name_without_ext)
    if len(total_text) < 2000:
        record = {
            "id": f"case_{safe_name}",
            "hierarchy": {
                "book": "",
                "subbook": "",
                "chapter": "裁判文书",
                "section": ""
            },
            "content": total_text,
            "source": court_name,
            "doc_type": "case",
            "case_number": case_number,
            "court": court_name,
            "date": judgment_date
        }
        legal_records.append(record)
    else:
        for section_name, section_texts in sections.items():
            if not section_texts:
                continue
            content = "\n".join(section_texts)
            if len(content) < 20:
                continue
            record = {
                "id": f"case_{safe_name}_{section_name}",
                "article_number": case_number,
                "hierarchy": {
                    "book": "",
                    "subbook": "",
                    "chapter": section_name,
                    "section": ""
                },
                "content": content,
                "source": court_name,
                "doc_type": "case",
                "case_number": case_number,
                "court": court_name,
                "date": judgment_date
            }
            legal_records.append(record)

    output_dir = os.path.dirname(output_json)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(legal_records, f, ensure_ascii=False, indent=4)

    print(f"案例解析结束！提取 {len(legal_records)} 条记录。")
    print(f"JSON已保存至：{os.path.abspath(output_json)}")
